from cryptography.fernet import Fernet
import os
from docx import Document

class FileEncryptor:
    def __init__(self, file_name):
        self.file_name = file_name
        self.key = Fernet.generate_key()
        self.cipher_suite = Fernet(self.key)
        self.encrypted_file_name = file_name + '.enc'

    def encrypt_file(self):
        # Read the file to encrypt
        with open(self.file_name, 'rb') as file:
            file_data = file.read()
        # Encrypt the file data
        encrypted_data = self.cipher_suite.encrypt(file_data)
        # Write the encrypted data to a new file
        with open(self.encrypted_file_name, 'wb') as file:
            file.write(encrypted_data)
        print(f"Encryption key: {self.key.decode()}")

    def delete_original_file(self):
        # Check if the file exists and remove it
        if os.path.exists(self.file_name):
            os.remove(self.file_name)
            print(f"Deleted {self.file_name}.")
        else:
            print(f"File '{self.file_name}' not found.")

    def save_key_to_docx(self, docx_file_path):
        # Load or create the DOCX document
        if os.path.exists(docx_file_path):
            doc = Document(docx_file_path)
        else:
            doc = Document()  # Create a new document if it doesn't exist
            doc.add_paragraph('')  # Add a first paragraph
        
        # Ensure there is at least two lines, clear the second line if exists
        while len(doc.paragraphs) < 2:
            doc.add_paragraph('')  # Add enough paragraphs
        
        # Add title:
        doc.paragraphs[0].add_run('Bank Statement Encryption Key:')

        # Clear the second line if it has content
        doc.paragraphs[1].clear()
        
        # Add the encryption key to the second line
        doc.paragraphs[1].add_run(self.key.decode())

        # Save the document
        doc.save(docx_file_path)
        print(f"Encryption key saved to {docx_file_path}")

